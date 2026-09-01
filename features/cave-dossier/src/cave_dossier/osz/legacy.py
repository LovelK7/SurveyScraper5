"""Legacy OSZ parser — pre-v10 zapisnici, ported from crospeleo-automation.

Port of the extraction core of ``services/osz_parser.py`` (docs/PORTING.md):
the FieldSpec alias table (union of every template generation observed in
the archive — the 2019-era Gauss-Krüger layout and the 2025 "kamp" layout
both), the sticky-label paragraph pass, the table pass with stacked-value
rows and label-in-cell-N/value-in-cell-N+1 rows, and — critically — the
**bold-run selection detection**: in every legacy generation the choice
fields (Vrsta objekta, Hidrološka karakteristika, …) print the FULL option
list and mark the selection in bold, so a plain-text read loses them.
python-docx is the document model, exactly as in the original (legacy
documents carry no ``w:sdt`` controls, so it sees everything).

Deliberately NOT ported: person/organization registry resolution, photo
author reconciliation, dossier warnings — the migration keeps recorded text
verbatim. ADDED beyond crospeleo (it drops these; migration wants them):
specs for Katastarski broj, Broj pločice, Županija, Grad/općina, Broj
ulaza, Topografski snimili (→ Crtali) and Mjerili.

The output is ``parse_legacy_osz(path) -> LegacyContent`` and
``to_v10_fields(...)`` which maps canonical keys onto the v10 address-map
keys plus checkbox tick candidates for the writer.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from cave_dossier.core.normalization import (
    cleanup_whitespace,
    normalize_lookup_key,
    parse_optional_float,
)


class LegacyParseError(RuntimeError):
    """The file could not be read as a legacy OSZ; message is CLI-ready."""


# ── FieldSpec table (verbatim from crospeleo _BASE_FIELD_SPECS) ──────
@dataclass(frozen=True)
class FieldSpec:
    canonical_name: str
    aliases: tuple[str, ...]
    required: bool = False


_BASE_FIELD_SPECS = (
    FieldSpec("object_name", ("ime objekta", "naziv objekta"), required=True),
    FieldSpec("synonyms", ("sinonimi", "sinonim")),
    FieldSpec("origin_of_name", ("podrijetlo imena", "porijeklo imena", "podrijetlo naziva")),
    FieldSpec("nearest_place", ("najblize mjesto", "najbliže mjesto", "najblize naselje", "najbliže naselje"), required=True),
    FieldSpec("locality", ("lokalitet", "lokaliteti"), required=True),
    FieldSpec(
        "location_access_text",
        ("polozaj i pristup", "položaj i pristup", "lokacija i pristup",
         "polozaj i pristup objektu", "položaj i pristup objektu"),
        required=True,
    ),
    FieldSpec("entrance_elevation_m", ("kota ulaza [m]", "kota ulaza", "kota ulaza (m.n.v.)")),
    FieldSpec("length_m", ("duljina", "duzina", "stvarna duljina [m]", "stvarna duljina")),
    FieldSpec("depth_m", ("dubina", "dubina [m]")),
    FieldSpec(
        "horizontal_length_m",
        ("horizontalna duljina", "horizontalna duzina", "horizont. dulj. [m]",
         "horizont. dulj.", "horizontalna dulj. [m]", "horizontalna duljina (m)"),
    ),
    FieldSpec(
        "vertical_difference_m",
        ("vertikalna razlika", "vertikalna razlika u visini", "visinska razlika [m]", "visinska razlika"),
    ),
    FieldSpec("coordinate_source", ("odredene po", "određene po", "izvor koordinata")),
    FieldSpec("entrance_dimensions",
              ("dimenzije ulaza [m x m]", "dimenzije ulaza", "dimenzije ulaza [mxm]",
               "dimenzije glavnog ulaza (m x m)", "dimenzije glavnog ulaza")),
    FieldSpec("cave_entrance_width_height",
              ("spiljski ulaz sirina visina", "špiljski ulaz širina visina")),
    FieldSpec("entrance_vertical_m", ("ulazna vertikala [m]", "ulazna vertikala")),
    FieldSpec("object_type", ("tip objekta", "vrsta objekta"), required=True),
    FieldSpec("morphological_type", ("morfoloski tip", "morfološki tip")),
    FieldSpec("hydrogeological_function", ("hidrogeoloska funkcija", "hidrogeološka funkcija")),
    FieldSpec("hydrological_characteristic", ("hidroloska karakteristika", "hidrološka karakteristika")),
    FieldSpec("technical_description",
              ("tehnicki opis", "tehnički opis", "osnovni opis s tehnickim podacima",
               "osnovni opis s tehničkim podacima"), required=True),
    FieldSpec(
        "research_period",
        ("razdoblje istrazivanja", "razdoblje istraživanja", "period istrazivanja",
         "period istraživanja", "vrijeme/datum istrazivanja", "vrijeme/datum istraživanja",
         "razdoblje ili datum istrazivanja", "razdoblje ili datum istraživanja"),
        required=True,
    ),
    FieldSpec("team_members", ("clanovi ekipe", "članovi ekipe", "sudionici"), required=True),
    FieldSpec(
        "organizations",
        ("organizacije koje su istrazivale", "organizacije koje su istraživale",
         "organizacije", "udruge", "istrazili (organizacije)", "istražili (organizacije)",
         "istrazivali (organizacije)", "istraživali (organizacije)"),
        required=True,
    ),
    FieldSpec(
        "recorder",
        ("zapisnicar", "zapisničar", "zapisnik sastavio", "zabiljezio", "zabilježio",
         "zapisnik ispunio/la", "zapisnik ispunio"),
        required=True,
    ),
    FieldSpec(
        "future_exploration_perspective",
        ("perspektiva buducih istrazivanja", "perspektiva budućih istraživanja",
         "buduca istrazivanja", "buduća istraživanja",
         "perspektiva daljnjeg istrazivanja", "perspektiva daljnjeg istraživanja"),
    ),
    FieldSpec(
        "expert_geology_hydrogeology",
        ("geološki i hidrogeološki", "geoloski i hidrogeoloski",
         "speleomorfološki geološki i hidrogeološki", "speleomorfoloski geoloski i hidrogeoloski"),
    ),
    FieldSpec("expert_meteorology", ("meteorološki", "meteoroloski", "mikroklimatski")),
    FieldSpec("expert_biology", ("biološki", "bioloski", "biospeleološki", "biospeleoloski")),
    FieldSpec("expert_archaeology_paleontology",
              ("arheološki i paleontološki", "arheoloski i paleontoloski")),
    FieldSpec("expert_pollution_anthropogenic",
              ("zagađenost i antropogeni utjecaji", "zagadjenost i antropogeni utjecaji")),
    FieldSpec("expert_hazards", ("opasnosti",)),
    FieldSpec("historical_data", ("povijesni podaci", "povijest istraživanja", "povijest istrazivanja")),
    FieldSpec("literature", ("literatura",)),
    FieldSpec("object_note", ("napomene", "napomena")),
    FieldSpec("photographed_by", ("fotografirali",)),
    FieldSpec("entrance_photo_author",
              ("autor fotografije ulaza", "autor fotografije", "fotografija ulaza")),
    # ── additions beyond crospeleo (it drops these; migration wants them) ──
    FieldSpec("cadastre_number", ("katastarski broj",)),
    FieldSpec("plaque_number", ("broj plocice", "broj pločice")),
    FieldSpec("county", ("zupanija", "županija")),
    FieldSpec("municipality", ("grad/opcina", "grad/općina", "grad opcina", "grad općina")),
    FieldSpec("entrance_count", ("broj ulaza",)),
    FieldSpec("surveyed_by", ("topografski snimili", "topo. snimili", "topo snimili", "crtali")),
    FieldSpec("measured_by", ("mjerili",)),
)

_FIELD_BY_ALIAS = {
    normalize_lookup_key(alias): spec
    for spec in _BASE_FIELD_SPECS
    for alias in spec.aliases
}

CHECKBOX_LIST_FIELDS = frozenset({
    "object_type", "morphological_type",
    "hydrological_characteristic", "hydrogeological_function",
})
_LIST_STYLE_FIELDS = frozenset({"organizations", "team_members", "synonyms"})
_LABEL_PREFIX_FALLBACKS = (
    ("polozajipristup", "location_access_text"),
    ("pristup", "location_access_text"),
)
_NO_CONTENT_SENTINELS = {"/", "nije poznato", "nepoznato"}
_LEADING_INDEX_PREFIX_RE = re.compile(r"^\s*\d+\s*[)\].\-:]\s*")
_LEADING_ORPHAN_COLON_RE = re.compile(r"^:\s*")

# Legacy two-part technical description (the 2019-era template).
_LEGACY_DESCRIPTION_PREFIXES = (
    "Opis neposredne okolice, ulaza i ulaznog dijela",
    "Opis unutrašnjosti",
)


@dataclass(frozen=True)
class LegacyContent:
    """Everything extracted from a legacy zapisnik."""

    fields: dict[str, str]              # canonical crospeleo-style keys
    bold_selections: dict[str, str]     # checkbox-list fields ("; "-joined)
    missing_required: tuple[str, ...] = field(default=())


# ── document ingestion ───────────────────────────────────────────────
def parse_legacy_osz(path: Path) -> LegacyContent:
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise LegacyParseError(
            "python-docx nije instaliran (extra [osz]) — legacy parser nedostupan."
        ) from exc
    if path.suffix.lower() != ".docx":
        raise LegacyParseError(f"{path.name}: legacy parser čita samo .docx "
                               "(stari .doc pretvori u Wordu prvo).")
    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001 — python-docx raises broadly
        raise LegacyParseError(f"{path.name}: nije čitljiv DOCX ({exc})") from exc

    parsed: dict[str, str] = {}
    _paragraph_pass(document, parsed)
    for table in document.tables:
        for key, value in _extract_table_fields(table).items():
            _merge_field(parsed, key, value)

    _normalize_parsed(parsed)
    missing = tuple(spec.canonical_name for spec in _BASE_FIELD_SPECS
                    if spec.required and not parsed.get(spec.canonical_name))
    bold = {key: parsed.pop(key) for key in list(parsed)
            if key in CHECKBOX_LIST_FIELDS}
    return LegacyContent(fields=parsed, bold_selections=bold, missing_required=missing)


def _paragraph_pass(document, parsed: dict[str, str]) -> None:
    current_key: str | None = None
    bold_section_key: str | None = None
    for paragraph in document.paragraphs:
        text = cleanup_whitespace(paragraph.text)
        if not text:
            continue
        maybe_pair = _split_key_value(text)
        if maybe_pair:
            key, value = maybe_pair
            current_key = key
            _merge_field(parsed, key, value)
            bold_section_key = current_key if value == "" else None
            continue
        if bold_section_key:
            bold_values = _extract_bold_values(paragraph, skip_key=bold_section_key)
            if bold_values:
                _merge_field(parsed, bold_section_key, "; ".join(bold_values))
                continue
        if current_key:
            _merge_field(parsed, current_key, text)
            bold_section_key = None


# ── label matching ───────────────────────────────────────────────────
def _canonical_key(raw_key: str, *, allow_prefix_fallback: bool = False) -> str | None:
    normalized = normalize_lookup_key(raw_key.rstrip(":"))
    if not normalized:
        return None
    spec = _FIELD_BY_ALIAS.get(normalized)
    if spec is not None:
        return spec.canonical_name
    if not allow_prefix_fallback:
        return None
    for prefix, canonical in _LABEL_PREFIX_FALLBACKS:
        if normalized.startswith(prefix):
            return canonical
    # Doubled section headers ("Položaj Položaj i pristup objektu:") —
    # rescue by alias suffix, long aliases only.
    for alias_norm, spec in _FIELD_BY_ALIAS.items():
        if len(alias_norm) >= 8 and normalized.endswith(alias_norm) and normalized != alias_norm:
            return spec.canonical_name
    return None


def _split_key_value(text: str) -> tuple[str, str] | None:
    if ":" not in text:
        return None
    raw_key, raw_value = text.split(":", 1)
    canonical = _canonical_key(raw_key, allow_prefix_fallback=True)
    if not canonical:
        return None
    return canonical, cleanup_whitespace(raw_value) or ""


def _merge_field(parsed: dict[str, str], key: str, value: str | None) -> None:
    cleaned = cleanup_whitespace(value)
    if not cleaned:
        return
    current = cleanup_whitespace(parsed.get(key))
    if current and cleaned != current:
        parsed[key] = f"{current}\n{cleaned}"
    elif not current:
        parsed[key] = cleaned


# ── table pass ───────────────────────────────────────────────────────
def _extract_table_fields(table) -> dict[str, str]:
    parsed: dict[str, str] = {}
    rows = table.rows
    for index, row in enumerate(rows):
        row_matches = _extract_row_matches(row)
        label_only = _extract_label_only_keys(row)
        if label_only and index + 1 < len(rows) and _row_has_no_keys(rows[index + 1]):
            if len(label_only) == 1:
                key = label_only[0]
                follow_idx = index + 1
                while follow_idx < len(rows) and _row_is_value_carrier(rows[follow_idx]):
                    value = _row_value_text(rows[follow_idx])
                    if value:
                        _append_carried_value(row_matches, key, value)
                    follow_idx += 1
            pair = _extract_columnar_checkbox_pair(row, rows[index + 1])
            for key, value in pair.items():
                row_matches.setdefault(key, value)
        for key, value in row_matches.items():
            _merge_field(parsed, key, value)
    return parsed


def _extract_row_matches(row) -> dict[str, str]:
    matches: dict[str, str] = {}
    segments = _row_segments(row)

    for canonical in CHECKBOX_LIST_FIELDS:
        bold = _extract_bold_values_for_key(row, canonical)
        if bold:
            matches[canonical] = "; ".join(bold)

    for index, segment in enumerate(segments):
        pair = _split_key_value(segment)
        if pair:
            key, value = pair
            if key in CHECKBOX_LIST_FIELDS:
                continue  # cell text carries ALL options, not the selection
            if value:
                matches.setdefault(key, value)
                continue
            trailing = _join_trailing_segments(segments[index + 1:])
            if trailing:
                matches.setdefault(key, trailing)
            continue
        canonical = _canonical_key(segment, allow_prefix_fallback=True)
        if canonical:
            if canonical in CHECKBOX_LIST_FIELDS:
                continue
            trailing = _join_trailing_segments(segments[index + 1:])
            if trailing:
                matches.setdefault(canonical, trailing)

    row_text = " ".join(segments)
    if row_text:
        for prefix in _LEGACY_DESCRIPTION_PREFIXES:
            if normalize_lookup_key(row_text).startswith(normalize_lookup_key(prefix)):
                _merge_field(matches, "technical_description", _strip_label_prefix(row_text))
                break
    _extract_signature_row(segments, row_text, matches)
    return matches


# "U | Kastvu | , dne | 15. 06. 2025. | Zapisnik ispunio/la: | …" — the
# place and date of filling out the form. No crospeleo counterpart (their
# parser drops both; user asked for them 2026-08-31). The 2019 generation
# omits the "U" cell ("Postojna | , dne | 12.03.2019. | …").
_SIGNATURE_INLINE_RE = re.compile(
    r"(?:\bU\s+)?(?P<place>[^,|]{2,40}?)\s*,?\s*\bdne\b\s*(?P<date>[\d][\d. ]*\d\.?)",
    re.IGNORECASE,
)


def _extract_signature_row(segments: list[str], row_text: str,
                           matches: dict[str, str]) -> None:
    if "record_place" in matches or not re.search(r"\bdne\b", row_text, re.IGNORECASE):
        return
    dne_index = next(
        (i for i, s in enumerate(segments)
         if normalize_lookup_key(s) in ("dne", "udne")), None,
    )
    if dne_index is not None:
        # Segment layout: place before "dne" (skipping a bare "U"), the
        # first digit-bearing segment after it is the date.
        place = next(
            (segments[i] for i in range(dne_index - 1, -1, -1)
             if normalize_lookup_key(segments[i]) not in ("u", "") and ":" not in segments[i]),
            None,
        )
        date = next(
            (s for s in segments[dne_index + 1:] if re.search(r"\d", s)), None,
        )
        if place:
            matches.setdefault("record_place", place.strip(" ,"))
        if date:
            matches.setdefault("record_date", date.strip())
        if place or date:
            return
    # Single-cell variant: "U Kastvu, dne 15.06.2025. Zapisnik ispunio…"
    inline = _SIGNATURE_INLINE_RE.search(row_text)
    if inline:
        matches.setdefault("record_place", inline.group("place").strip(" ,"))
        matches.setdefault("record_date", inline.group("date").strip())


def _row_segments(row) -> list[str]:
    segments: list[str] = []
    for cell in row.cells:
        text = cleanup_whitespace(cell.text)
        if text and (not segments or segments[-1] != text):
            segments.append(text)
    return segments


def _join_trailing_segments(segments: list[str]) -> str:
    """Value cells after a label cell, stopping at the next recognised
    label (label-with-value pairs like 'Izvor kote ulaza:' that carry NO
    spec keep flowing — _split_entrance_elevation_parts relies on it)."""
    collected: list[str] = []
    for segment in segments:
        if _canonical_key(segment.split(":", 1)[0]) is not None:
            break
        collected.append(segment)
    return cleanup_whitespace(" ".join(collected)) or ""


def _extract_label_only_keys(row) -> list[str]:
    keys: list[str] = []
    for segment in _row_segments(row):
        canonical = _canonical_key(segment, allow_prefix_fallback=True)
        if canonical and canonical not in keys:
            keys.append(canonical)
    return keys


def _row_has_no_keys(row) -> bool:
    return not _extract_label_only_keys(row)


def _row_is_value_carrier(row) -> bool:
    segments = _row_segments(row)
    if not segments:
        return False
    for segment in segments:
        if _canonical_key(segment, allow_prefix_fallback=True):
            return False
        stripped = _LEADING_ORPHAN_COLON_RE.sub(
            "", _LEADING_INDEX_PREFIX_RE.sub("", segment))
        colon = stripped.find(":")
        if colon >= 0 and not re.search(r"[.!?]", stripped[:colon]):
            return False
    return True


def _row_value_text(row) -> str:
    return cleanup_whitespace(" ".join(_row_segments(row))) or ""


def _append_carried_value(matches: dict[str, str], key: str, value: str) -> None:
    separator = "; " if key in _LIST_STYLE_FIELDS else "\n"
    current = matches.get(key)
    matches[key] = f"{current}{separator}{value}" if current else value


def _strip_label_prefix(text: str) -> str:
    if ":" in text:
        return cleanup_whitespace(text.split(":", 1)[1]) or ""
    return text


# ── bold-run selection detection ─────────────────────────────────────
def _extract_bold_values(paragraph, skip_key: str | None = None) -> list[str]:
    """Contiguous bold runs, joined with NO separator (Word splits 'špilja'
    into 'š'+'pilja' at the diacritic); whitespace-only bold runs preserve
    word boundaries."""
    values: list[str] = []
    buffer: list[str] = []

    def flush() -> None:
        text = cleanup_whitespace("".join(buffer))
        buffer.clear()
        if not text:
            return
        if skip_key and _canonical_key(text) == skip_key:
            return
        values.append(text)

    for run in paragraph.runs:
        if run.bold:
            buffer.append(run.text)
        else:
            flush()
    flush()
    return values


def _cell_keys(cell) -> set[str]:
    keys: set[str] = set()
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            canonical = _canonical_key(run.text)
            if canonical:
                keys.add(canonical)
        first_line = cleanup_whitespace(paragraph.text)
        if first_line:
            canonical = _canonical_key(first_line.split(":", 1)[0])
            if canonical:
                keys.add(canonical)
    return keys


def _extract_bold_values_for_key(row, canonical_key: str) -> list[str]:
    """Bold values scoped to the column that declares canonical_key."""
    collecting = False
    values: list[str] = []
    seen = set()
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        keys = _cell_keys(cell)
        if keys and canonical_key not in keys and collecting:
            break
        if canonical_key in keys:
            collecting = True
        if collecting:
            for paragraph in cell.paragraphs:
                values.extend(_extract_bold_values(paragraph, skip_key=canonical_key))
    return values


def _extract_columnar_checkbox_pair(header_row, value_row) -> dict[str, str]:
    """Newer legacy layout: header row 'Morfološki tip: | Hidrološka …'
    with the option lists (bold selection) in the NEXT row, same columns."""
    matches: dict[str, str] = {}
    header_cells = list(header_row.cells)
    value_cells = list(value_row.cells)
    for index, header in enumerate(header_cells):
        if index >= len(value_cells):
            break
        canonical = _canonical_key(cleanup_whitespace(header.text) or "",
                                   allow_prefix_fallback=False)
        if canonical not in CHECKBOX_LIST_FIELDS:
            continue
        bold: list[str] = []
        for paragraph in value_cells[index].paragraphs:
            bold.extend(_extract_bold_values(paragraph))
        if bold:
            matches[canonical] = "; ".join(dict.fromkeys(bold))
    return matches


# ── post-processing ──────────────────────────────────────────────────
def _normalize_parsed(parsed: dict[str, str]) -> None:
    elevation, source = _split_entrance_elevation_parts(parsed.get("entrance_elevation_m"))
    if elevation is not None:
        parsed["entrance_elevation_m"] = elevation
    if source:
        parsed["entrance_elevation_source"] = source
    for key in ("historical_data", "literature", "object_note"):
        if key in parsed and _is_no_content(parsed[key]):
            parsed.pop(key)
    note = parsed.get("object_note")
    if note and re.search(r"zapisni(k ispunio|čar|car)", note, re.IGNORECASE):
        parsed.pop("object_note")


def _split_entrance_elevation_parts(value: str | None) -> tuple[str | None, str | None]:
    cleaned = cleanup_whitespace(value)
    if not cleaned:
        return None, None
    patterns = (
        r"(?i)\bIzvor\s+kote\s+ulaza\s*:\s*",
        r"(?i)\bOdr\.?\s*po\s*:\s*",
    )
    for pattern in patterns:
        parts = re.split(pattern, cleaned, maxsplit=1)
        if len(parts) == 2:
            return cleanup_whitespace(parts[0]), cleanup_whitespace(parts[1])
    return cleaned, None


def _is_no_content(value: str) -> bool:
    return value.strip().rstrip(".,;").lower() in _NO_CONTENT_SENTINELS


# The CroSpeleo Izvor koordinata vocabulary (crospeleo _COORD_SOURCE_OPTIONS)
# — the legacy cell often drags neighbouring table junk along (the WGS84
# digit grid), so the recognised option is fished out of the raw text.
_COORD_SOURCE_OPTIONS = (
    "GPS", "HOK 1:5000", "LIDAR", "TK 1:100000", "TK 1:25000",
    "geodetski određene", "karta nepoznatog mjerila", "referenca",
)


def _normalize_coordinate_source(raw: str | None) -> str | None:
    cleaned = cleanup_whitespace(raw)
    if not cleaned:
        return None
    normalized = normalize_lookup_key(cleaned)
    for option in _COORD_SOURCE_OPTIONS:
        if normalize_lookup_key(option) == normalized:
            return option
    for option in _COORD_SOURCE_OPTIONS:
        option_norm = normalize_lookup_key(option)
        if option_norm in normalized or normalized in option_norm:
            return option
    return cleaned  # unrecognised — keep the recorder's text verbatim


def _extract_dimension_pair(value: str | None) -> tuple[float | None, float | None]:
    cleaned = cleanup_whitespace(value)
    if not cleaned or cleaned == "/":
        return None, None
    numbers = re.findall(r"-?\d+(?:[.,]\d+)?", cleaned)
    if len(numbers) < 2:
        return None, None
    first, second = (float(n.replace(",", ".")) for n in numbers[:2])
    return first, second


# ── legacy canonical keys -> v10 field keys ──────────────────────────
_DIRECT_V10 = {
    "object_name": "ime_objekta",
    "synonyms": "sinonimi",
    "cadastre_number": "katastarski_broj",
    "plaque_number": "broj_plocice",
    "county": "zupanija",
    "municipality": "grad_opcina",
    "nearest_place": "najblize_mjesto",
    "locality": "lokalitet",
    "location_access_text": "polozaj_pristup",
    "entrance_elevation_m": "kota_ulaza",
    "entrance_elevation_source": "izvor_kote",
    "coordinate_source": "izvor_koordinata",
    "entrance_count": "broj_ulaza",
    "technical_description": "opis",
    "future_exploration_perspective": "perspektiva",
    "expert_geology_hydrogeology": "geologija",
    "expert_meteorology": "mikroklima",
    "expert_biology": "biospeleologija",
    "expert_archaeology_paleontology": "arheologija",
    "expert_pollution_anthropogenic": "zagadenost",
    "expert_hazards": "opasnosti",
    "historical_data": "povijest",
    "literature": "literatura",
    "object_note": "napomene",
    "research_period": "datum_istrazivanja",
    "organizations": "istrazile_udruge",
    "team_members": "clanovi_ekipe",
    "surveyed_by": "crtali",
    "measured_by": "mjerili",
    "photographed_by": "fotografirali",
    "entrance_photo_author": "autor_fotografije_ulaza",
    "recorder": "zapisnicar",
    "record_place": "mjesto_zapisnika",
    "record_date": "datum_zapisnika",
}
_NUMERIC_V10 = {
    "length_m": "duljina",
    "depth_m": "dubina",
    "horizontal_length_m": "horizontalna_duljina",
    "vertical_difference_m": "visinska_razlika",
}


def to_v10_fields(content: LegacyContent) -> tuple[dict[str, str], list[str], list[str]]:
    """(v10 field values, checkbox tick candidates, notes).

    Values keep the recorder's text; numeric cells drop their unit suffix
    ("7 m" → "7"). Entrance size prefers the explicit width/height pair
    over "Dimenzije glavnog ulaza" (crospeleo RULES §4). Bold selections
    become tick candidates the writer matches against v10 labels —
    Morfološki tip has no v10 group and will surface as unmatched.
    """
    fields: dict[str, str] = {}
    notes: list[str] = []

    for legacy_key, v10_key in _DIRECT_V10.items():
        value = content.fields.get(legacy_key)
        if value:
            fields[v10_key] = value
    source = _normalize_coordinate_source(content.fields.get("coordinate_source"))
    if source:
        fields["izvor_koordinata"] = source
    for legacy_key, v10_key in _NUMERIC_V10.items():
        number = parse_optional_float(content.fields.get(legacy_key))
        if number is not None:
            fields[v10_key] = _number_text(number)

    if fields.get("autor_fotografije_ulaza", "").strip().lower() in ("da", "ne", "yes", "no"):
        fields.pop("autor_fotografije_ulaza")  # legacy yes/no flag, not a name

    width, height = _extract_dimension_pair(content.fields.get("cave_entrance_width_height"))
    if width is None:
        width, height = _extract_dimension_pair(content.fields.get("entrance_dimensions"))
    if width is not None:
        fields["sirina_ulaza"] = _number_text(width)
    if height is not None:
        fields["visina_duljina_ulaza"] = _number_text(height)

    vertical = parse_optional_float(content.fields.get("entrance_vertical_m"))
    if vertical is not None:
        # v10 dropped "Ulazna vertikala" — keep the fact in Napomene.
        addition = f"Ulazna vertikala: {_number_text(vertical)} m."
        fields["napomene"] = (f"{fields['napomene']}\n{addition}"
                              if fields.get("napomene") else addition)

    ticks: list[str] = []
    for legacy_key, joined in content.bold_selections.items():
        for value in joined.split(";"):
            value = value.strip()
            if not value:
                continue
            if legacy_key == "morphological_type":
                notes.append(f"Morfološki tip '{value}' nema v10 kućicu — provjeri ručno.")
                continue
            ticks.append(value)
    origin = content.fields.get("origin_of_name")
    if origin:
        ticks.append(origin)

    for missing in content.missing_required:
        notes.append(f"Stari OSZ nema obavezno polje: {missing}.")
    return fields, ticks, notes


def _number_text(value: float) -> str:
    if float(value).is_integer():
        return str(int(value))
    return f"{value:g}".replace(".", ",")
